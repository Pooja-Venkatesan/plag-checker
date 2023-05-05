#!/usr/bin/python3
from django.shortcuts import render, redirect
from django.urls import reverse
from django.http import HttpResponse
from plagiarismchecker.algorithm import main
from plagiarismchecker.algorithm.main import tracker,totalPercent, uniquePercent, outputLink, text

from docx import *
from docx import Document
from plagiarismchecker.algorithm import fileSimilarity
import PyPDF2 

# from django.shortcuts import redirect
# from xhtml2pdf import pisa
# from io import BytesIO
# from django.template.loader import get_template
# from django.views import View

# Create your views here.
#home

def home(request):
    return render(request, 'pc/index.html')

# def home(request):
#     if None in request.POST:
#            return render(request, 'pc/index.html')
    
#     if 'report-gen' in request.POST:
                
#         return render(request, 'pc/report.html')
# totalPercent, uniquePercent, outputLink, text, tracker= main.findSimilarity(request.POST['q'])

# ////////////////////////////////////////////////////////////
# def report(request):
#     # links = list(outputLink.keys())
#     # scores = [i for i in outputLink.values()]
#     if request.POST['q']: 
#         totalPercent, uniquePercent, links, scores, text, tracker= main.findSimilarity(request.POST['q'])
#         print("\n list of links:", links)
#         return render(request, 'pc/report.html', {'tracker' : tracker, 'links' : links, 'Scores':scores, 'text':text,'totalPercent':totalPercent, 'uniquePercent':uniquePercent})
# //////////////////////////
totalPercent = 0
uniquePercent = 0
links = []
scores = [] 
text = str 
tracker = {}


def report(request):
    if request.POST.get('q'): 
        totalPercent, uniquePercent, links, scores, text, tracker= main.findSimilarity(request.POST['q'])
        print("\n List of links:", links, "and Scores", scores )
        return render(request, 'pc/report.html', {'tracker' : tracker, 'links' : links, 'Scores':scores, 'text':text,'totalPercent':totalPercent, 'uniquePercent':uniquePercent})
    elif request.FILES.get(['docfile']):
        totalPercent, uniquePercent, links, scores, text, tracker= main.findSimilarity(value)
        return render(request, 'pc/index.html',{'links': links, 'totalPercent': totalPercent,'uniquePercent':uniquePercent, 'text' : text,'scores;':scores, 'tracker':tracker})
    else:
        return HttpResponse("No content to check for plagiarism.")



# def render_to_pdf(template_src, context_dict={}):
# 	template = get_template(template_src)
# 	html = template.render(context_dict)
# 	result = BytesIO()
# 	pdf = pisa.pisaDocument(BytesIO(html.encode("utf8")), result)
# 	if not pdf.err:
# 		return HttpResponse(result.getvalue(), content_type='application/pdf')
# 	return None

# class ViewPDF(View):
#     def get(self, request, *args, **kwargs):
#         if request.session.get("check-plag"):
#             data=request.session.get("check-plag")
#         pdf = render_to_pdf('pc/report.html', data)
#         return HttpResponse(pdf, content_type='application/pdf')

#web search(Text)

def test(request):
    # print("request is welcome test")

    print(request.POST['q'])  
    if request.POST['q']: 
        totalPercent, uniquePercent, links, scores, text, tracker= main.findSimilarity(request.POST['q'])
        uniquePercent = 100 - totalPercent
        totalPercent = round(totalPercent,2)
        uniquePercent = round(uniquePercent,2)
    print("\nOutput..!!!",totalPercent, uniquePercent, links,scores, text, tracker)
    print('\ntext:', text)
    print('\nlinks:', links)
    return render(request, 'pc/index.html',{'totalPercent': totalPercent,'uniquePercent': uniquePercent,'links': links,'scores': scores, 'text' : text, 'tracker':tracker})

    

#web search file(.txt, .docx, .pdf)
# def filetest(request):
#     value = '' 
#     print("GET QUERY STARTED!!!\n")  
#     print("Request:\n",request.FILES['docfile'])
#     if str(request.FILES['docfile']).endswith(".txt"):
#         value = str(request.FILES['docfile'].read())
#         value = value.decode()
#     elif str(request.FILES['docfile']).endswith(".docx"):
#         document = Document(request.FILES['docfile'])
#         for para in document.paragraphs:
#             value += para.text
#             value = value.decode()
#     elif str(request.FILES['docfile']).endswith(".pdf"):
#         # creating a pdf file object 
#         pdfFileObj = request.FILES['docfile']

#         # creating a pdf reader object 
#         pdfReader = PyPDF2.PdfFileReader(pdfFileObj)

#         # print number of pages in the pdf file
#         print("Number of pages:", pdfReader.getNumPages())

#         # creating a page object 
#         pageObj = pdfReader.getPage(0) 

#         # extract text from page
#         value = pageObj.extractText()
        
#         # closing the pdf file object 
#         pdfFileObj.close()
    
#     text = value
   

# # File input word limit set 
#     if len(value.split()) > 1500:
#         return render(request, 'pc/index.html', {'error': 'Word count limit exceeded (1500 words)'})
    
#     print("\n  TEXT IN FILE INPUT :",text)
#     totalPercent, uniquePercent, links, scores, text, tracker= main.findSimilarity(text)
#     text = value
#     text = text.replace('\r\n', '<br>')
#     text = text.replace("\xe2\x80\x99","'")
#     print("Output..!!! \n","\n totalPercent, uniquePercent, links, scores, text, tracker\n",totalPercent, uniquePercent, links, scores, text, tracker )
#     print("\n list of links:", links)
#     print("\n scores:", scores)
#     return render(request, 'pc/index.html',{'links': links, 'totalPercent': totalPercent,'uniquePercent':uniquePercent, 'text' : text,'scores;':scores, 'tracker':tracker})
def filetest(request):
    value = '' 
    print("GET QUERY STARTED!!!\n")  
    print("Request:\n", request.FILES['docfile'])

    file = request.FILES['docfile']
    # Check the file size
    if file.size > (2 * 1024 * 1024):  # 2MB limit (adjust the limit as needed)
        print("...........ERROR:File size limit exceeded (2MB).....................")
        return render(request, 'pc/index.html', {'error': 'File size limit exceeded (2MB)'})

    if str(file).endswith(".txt"):
        value = file.read()
        print('txt value:', value)
        value = value.decode('utf-8')
        print('after decode - txt value:', value)
    elif str(file).endswith(".docx"):
        document = Document(file)
        for para in document.paragraphs:
            value += para.text
        print('doc-value:', value)
        if isinstance(value, bytes):
            value = value.decode('utf-8')
            print('after decode - doc value:', value)
    elif str(file).endswith(".pdf"):
        # creating a pdf file object 
        pdfFileObj = file

        # creating a pdf reader object 
        pdfReader = PyPDF2.PdfFileReader(pdfFileObj)

        # print number of pages in the pdf file
        print("Number of pages:", pdfReader.getNumPages())

        # creating a page object 
        pageObj = pdfReader.getPage(0) 

        # extract text from page
        value = pageObj.extractText()
        print('pdf-value:', value)

        # closing the pdf file object 
        pdfFileObj.close()

    text = value

    # File input word limit set 
    if len(value.split()) > 1500:
        print("............ERROR: Word count limit exceeded (1500 words).............")
        return render(request, 'pc/index.html', {'error': 'Word count limit exceeded (1500 words)'})

    print("\n  TEXT IN FILE INPUT :",text)
    totalPercent, uniquePercent, links, scores, text, tracker= main.findSimilarity(text)
    text = str(value)
    text = text.replace('\r\n', '<br>')
    text = text.replace("\xe2\x80\x99","'")
    print("Output..!!! \n","\n totalPercent, uniquePercent, links, scores, text, tracker\n",totalPercent, uniquePercent, links, scores, text, tracker )
    print("\n list of links:", links)
    print("\n scores:", scores)
    
    return render(request, 'pc/index.html',{'links': links, 'totalPercent': totalPercent,'uniquePercent':uniquePercent, 'text' : text,'scores;':scores, 'tracker':tracker})


# def filetest(request):
#     value = ''
#     if 'docfile' in request.FILES:
#         if str(request.FILES['docfile']).endswith(".txt"):
#             value = str(request.FILES['docfile'].read())
#         elif str(request.FILES['docfile']).endswith(".docx"):
#             document = Document(request.FILES['docfile'])
#             for para in document.paragraphs:
#                 value += para.text
#         elif str(request.FILES['docfile']).endswith(".pdf"):
#             pdfFileObj = request.FILES['docfile']
#             pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
#             pageObj = pdfReader.getPage(0)
#             value = pageObj.extractText()
#             text = value
#             pdfFileObj.close()

#             totalPercent, uniquePercent, links, scores, text, tracker= main.findSimilarity(value)
#             links = list(outputLink.keys())
#             scores = list(outputLink.values())
#             print("\n list of links:", links)
#             print("\n scores:", scores)
#             return render(request, 'pc/index.html',{'links': links, 'totalPercent': totalPercent,'uniquePercent':uniquePercent, 'text' : text,'scores':scores, 'tracker':tracker})
#     else:
#         return HttpResponse("No file uploaded.")        


# .........Extrinsic.................
#text compare
def fileCompare(request):
    return render(request, 'pc/doc_compare.html') 

#two text compare(Text)
def twofiletest1(request):
    print("Submiited text for 1st and 2nd")
    print(request.POST['q1'])
    print(request.POST['q2'])

    if request.POST['q1'] != '' and request.POST['q2'] != '': 
        print("Got both the texts")
        result = fileSimilarity.findFileSimilarity(request.POST['q1'],request.POST['q2'])
    result = round(result,2)    
    print("Output>>>!!!!",result)
    return render(request, 'pc/doc_compare.html',{'result': result})
    

#two text compare(.txt, .docx)
def twofilecompare1(request):
    value1 = ''
    value2 = ''
    if (str(request.FILES['docfile1'])).endswith(".txt") and (str(request.FILES['docfile2'])).endswith(".txt"):
        value1 = str(request.FILES['docfile1'].read())
        value2 = str(request.FILES['docfile2'].read())

    elif (str(request.FILES['docfile1'])).endswith(".docx") and (str(request.FILES['docfile2'])).endswith(".docx"):
        document = Document(request.FILES['docfile1'])
        for para in document.paragraphs:
            value1 += para.text
        document = Document(request.FILES['docfile2'])
        for para in document.paragraphs:
            value2 += para.text

    result = fileSimilarity.findFileSimilarity(value1,value2)
    
    print("Output! \n",result)
    return render(request, 'pc/doc_compare.html',{'result': result})
# .................
